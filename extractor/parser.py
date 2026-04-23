import os
import re
import json
import logging
from dataclasses import dataclass, field
from typing import TypedDict

logger = logging.getLogger(__name__)

HTML_TAG_RE = re.compile(r'<[a-zA-Z][^>]*>')


class Clause(TypedDict):
    number: str
    content: str
    level: int
    path: str
    full_name: str


@dataclass
class ParsedLine:
    text: str
    is_heading_style: bool = False
    heading_level: int = 0  # 0 = 非标题；1/2/3... = Word Heading 层级


def _table_to_narrative(table) -> str:
    """
    将 docx 表格转为口语化的文字描述。
    以表头为 key，逐行将每个单元格内容以"字段名：xxx"的方式串联。
    """
    rows = table.rows
    if len(rows) == 0:
        return ""

    headers = [cell.text.strip().replace('\n', ' ') for cell in rows[0].cells]
    seen_headers = []
    seen_indices = []
    for i, h in enumerate(headers):
        if h not in seen_headers:
            seen_headers.append(h)
            seen_indices.append(i)

    if len(rows) == 1:
        return "【表格】表头：" + "、".join(seen_headers)

    narrative_lines = ["【表格内容】"]
    for row_idx, row in enumerate(rows[1:], start=1):
        cells = [cell.text.strip().replace('\n', '；') for cell in row.cells]
        parts = []
        for i, col_idx in enumerate(seen_indices):
            header = seen_headers[i]
            value = cells[col_idx] if col_idx < len(cells) else ""
            if value:
                parts.append(f"{header}：{value}")
        if parts:
            narrative_lines.append(f"  第{row_idx}项：{'；'.join(parts)}。")

    return "\n".join(narrative_lines)


def _is_heading_style(style_name: str | None) -> bool:
    """判断段落样式是否为 Heading 样式"""
    if not style_name:
        return False
    name = style_name.lower()
    if name.startswith('heading') or name.startswith('标题'):
        return True
    if name.startswith('toc'):
        return False
    return False


def _get_heading_level(style_name: str | None) -> int:
    """从 Word 标题样式名提取层级数字，例如 'Heading 1' / '标题1' → 1；非标题 → 0"""
    if not style_name:
        return 0
    name = style_name.lower().strip()
    if name.startswith('heading') or name.startswith('标题'):
        m = re.search(r'(\d+)\s*$', name)
        if m:
            return int(m.group(1))
    return 0


def parse_docx(filepath: str) -> list[ParsedLine]:
    """
    解析 docx 文件，返回结构化行数据。
    每行标记是否来自 Heading 样式的段落，供下游做两层识别：
      1) 有 Heading 样式 → 再检查文本是否匹配 1./1.1./1.1.1. 格式
      2) 无 Heading 样式 → 视为普通内容
    表格以口语化方式转译。
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(filepath)
    body = doc.element.body

    table_elements = [tbl._element for tbl in doc.tables]
    table_map = {id(el): tbl for el, tbl in zip(table_elements, doc.tables)}

    para_map = {}
    for para in doc.paragraphs:
        para_map[id(para._element)] = para

    result: list[ParsedLine] = []

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            texts = [node.text for node in child.iter(qn('w:t')) if node.text]
            para_text = ''.join(texts)

            para_obj = para_map.get(id(child))
            heading = False
            hlevel = 0
            if para_obj is not None:
                style_name = para_obj.style.name if para_obj.style else None
                heading = _is_heading_style(style_name)
                hlevel = _get_heading_level(style_name)

            result.append(ParsedLine(text=para_text, is_heading_style=heading, heading_level=hlevel))

        elif tag == 'tbl':
            tbl_obj = table_map.get(id(child))
            if tbl_obj is None:
                tbl_obj = Table(child, doc)
            narrative = _table_to_narrative(tbl_obj)
            if narrative:
                for line in narrative.split('\n'):
                    result.append(ParsedLine(text=line, is_heading_style=False))

    return result


def parse_txt(filepath: str) -> list[ParsedLine]:
    """解析 txt 文件，自动检测编码（utf-8 优先，gbk 兜底）。无样式信息。"""
    for encoding in ("utf-8", "gbk", "utf-8-sig", "latin-1"):
        try:
            with open(filepath, "r", encoding=encoding) as f:
                content = f.read()
            return [ParsedLine(text=line, is_heading_style=False) for line in content.split('\n')]
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文件编码: {filepath}")


def _expand_html_table_grid(table_tag) -> list[list[str]]:
    """将 HTML 表格展开为二维文本 grid，正确处理 rowspan/colspan。

    合并单元格的文本会被复制填充到其占据的每一个细单元格中，
    确保展开后的 grid 每行列数一致、数据完整。
    """
    rows = table_tag.find_all('tr')
    if not rows:
        return []

    max_cols = 0
    for row in rows:
        cols_in_row = sum(int(c.get('colspan', 1)) for c in row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols_in_row)

    grid: list[list[str]] = []
    fill: dict[tuple[int, int], str] = {}

    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        grid_row = [''] * max_cols
        cell_idx = 0
        col_pos = 0

        while col_pos < max_cols:
            if (row_idx, col_pos) in fill:
                grid_row[col_pos] = fill[(row_idx, col_pos)]
                col_pos += 1
                continue

            if cell_idx >= len(cells):
                col_pos += 1
                continue

            cell = cells[cell_idx]
            text = cell.get_text(strip=True).replace('\n', '；')
            rs = int(cell.get('rowspan', 1))
            cs = int(cell.get('colspan', 1))

            for c in range(cs):
                grid_row[col_pos + c] = text
            for r in range(1, rs):
                for c in range(cs):
                    fill[(row_idx + r, col_pos + c)] = text

            cell_idx += 1
            col_pos += cs

        grid.append(grid_row)

    return grid


def _build_narrative_headers(
    raw_header_row: list[str],
    is_header_row: bool = True,
) -> tuple[list[str], list[int]]:
    """从展开后的表头行构建去重的列名列表。

    - is_header_row=True（默认，表头场景）：
      相邻位置出现的同文本（即 _expand_html_table_grid 在 colspan 处复制出的副本）
      被视作同一列，只保留第一个出现位置；非相邻同名仍按"列名1、列名2…"区分。
      例：['描述','描述','价格','描述','描述']
          -> 头2个相邻折叠成一份'描述'，尾2个相邻折叠成一份'描述'，
             两段非相邻 → ('描述1', '价格', '描述2')。
    - is_header_row=False（数据行场景）：
      不做相邻折叠，所有非空同名一律加 1/2/3 后缀（保留旧行为）。

    返回 (seen_headers, seen_indices)，seen_indices 是这些列在原 grid 中的列下标。
    """
    headers: list[str] = [re.sub(r'\s+', '', h) for h in raw_header_row]

    kept_positions: list[int] = []
    for i, h in enumerate(headers):
        if not h:
            continue
        if is_header_row and i > 0 and headers[i - 1] == h:
            continue
        kept_positions.append(i)

    name_positions: dict[str, list[int]] = {}
    for i in kept_positions:
        name_positions.setdefault(headers[i], []).append(i)

    final_names: dict[int, str] = {}
    for name, positions in name_positions.items():
        if len(positions) == 1:
            final_names[positions[0]] = name
        else:
            for seq, pos in enumerate(positions, start=1):
                final_names[pos] = f"{name}{seq}"

    seen_headers: list[str] = []
    seen_indices: list[int] = []
    for i in kept_positions:
        h = final_names.get(i, '')
        if h and h not in seen_headers:
            seen_headers.append(h)
            seen_indices.append(i)

    return seen_headers, seen_indices


def _html_table_to_narrative(table_tag) -> str:
    """将单个 HTML <table> 标签转为口语化编号列表，每条数据以（1）（2）… 列出。

    正确处理 rowspan/colspan：合并单元格的文本填充到所有占据的位置。
    表头若存在横向合并（colspan），展开为"表头名1、表头名2…"。

    前导行处理（兼容 Quill-better-table 等富文本编辑器产物）：
      - 全空行（仅含 <br> 等占位）：直接跳过；
      - 整行同文本行（即一个 colspan=全列宽 的合并单元格，常见于"表标题"）：
        收集为 caption，narrative 以「表标题：xxx」前缀输出；
      - 第一个既不是全空、也不是 caption 的行视为真正的表头。
    """
    grid = _expand_html_table_grid(table_tag)
    if not grid:
        return ''

    captions: list[str] = []
    header_idx = 0
    while header_idx < len(grid):
        row = grid[header_idx]
        non_empty = [c for c in row if c]
        if not non_empty:
            header_idx += 1
            continue
        if len(non_empty) == len(row) and len(set(non_empty)) == 1:
            captions.append(non_empty[0])
            header_idx += 1
            continue
        break

    if header_idx >= len(grid):
        return "表标题：" + "；".join(captions) if captions else ''

    seen_headers, seen_indices = _build_narrative_headers(grid[header_idx])

    if not seen_headers:
        return "表标题：" + "；".join(captions) if captions else ''

    data_rows = grid[header_idx + 1:]

    if not data_rows:
        body = "表头：" + "、".join(seen_headers)
    else:
        lines: list[str] = []
        for row_idx, row in enumerate(data_rows, start=1):
            parts: list[str] = []
            for i, col_idx in enumerate(seen_indices):
                header = seen_headers[i]
                value = row[col_idx] if col_idx < len(row) else ''
                if value:
                    parts.append(f"{header}：{value}")
            if parts:
                lines.append(f"（{row_idx}）{'；'.join(parts)}。")
        body = '\n'.join(lines)

    if captions:
        return "表标题：" + "；".join(captions) + "\n" + body
    return body


def extract_clause_references(html: str) -> list[dict]:
    """从 clauseContent HTML 中抽取带有 data-policy-id 的 <span> 作为引用。

    识别规则（与 class 无关，仅看 data-* 属性）：
      - 标签必须是 <span>
      - 必须有 data-policy-id 且 strip 后非空
      - data-clause-id 可缺失或为空串；空串语义保留为"引用整篇 policy"
        （由 resolve_references → _expand_to_clause_nodes 处理）

    返回结构示例（每条 ref 含 5 个字段，resolvedClauses/cycle 由后续 resolve_references 填充）::

        [
            {
                "policyId": "KH...",
                "clauseId": "id1,id2"  # 保留原串，可能是单 id / 逗号多 id / 空串
                "highlightedContent": "...",
                "resolvedClauses": [],
                "cycle": False,
            },
            ...
        ]

    - 没有 data-policy-id（或为空）的 span 直接跳过，不打 warning（普通 span 太常见，避免噪音）。
    - 同一 span 多次出现保留多次，按文档出现顺序。
    """
    if not html or '<span' not in html:
        return []

    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, 'html.parser')
    spans = soup.find_all('span', attrs={'data-policy-id': True})
    refs: list[dict] = []
    for span in spans:
        policy_id = (span.get('data-policy-id') or '').strip()
        if not policy_id:
            continue
        clause_id = (span.get('data-clause-id') or '').strip()
        highlighted = span.get_text()
        refs.append({
            'policyId': policy_id,
            'clauseId': clause_id,
            'highlightedContent': highlighted,
            'resolvedClauses': [],
            'cycle': False,
        })
    return refs


def _convert_html_to_markdown(text: str, narrativize: bool = True) -> str:
    """将文本中的 HTML 片段转为 Markdown，纯文本原样返回。

    Args:
        text: 原始文本（可能包含 HTML）
        narrativize: 若为 True，将 HTML 表格转为口语化编号列表（（1）（2）…），
                     结合字段名称描述每条数据，便于 LLM 理解；
                     若为 False，使用常规 Markdown 表格格式。
    """
    if not HTML_TAG_RE.search(text):
        return text

    from markdownify import markdownify as md

    if not narrativize:
        return md(text, strip=['span', 'div', 'colgroup', 'col']).strip()

    from bs4 import BeautifulSoup, NavigableString

    soup = BeautifulSoup(text, 'html.parser')
    tables = soup.find_all('table')

    if not tables:
        return md(text, strip=['span', 'div', 'colgroup', 'col']).strip()

    for table in tables:
        narrative = _html_table_to_narrative(table)
        table.replace_with(NavigableString(narrative))

    return md(str(soup), strip=['span', 'div', 'colgroup', 'col']).strip()


def _derive_parent_number(number: str) -> str:
    """
    从层级编号推导其直接父级编号。
    例如 '1.2.1' → '1.2'，'1.2' → '1'，'1' → ''，'前言' → ''
    """
    parts = number.split('.')
    if len(parts) <= 1:
        return ''
    return '.'.join(parts[:-1])


def _build_full_path(number: str, node_map: dict[str, dict]) -> str:
    """
    从 number 向上遍历父级链，构建完整祖先路径。
    返回格式如 '1/1.2/1.2.1'（不含自身编号）。
    """
    ancestors: list[str] = []
    current = number
    while True:
        parent_num = _derive_parent_number(current)
        if not parent_num or parent_num not in node_map:
            break
        ancestors.append(parent_num)
        current = parent_num
    ancestors.reverse()
    return '/'.join(ancestors)


def parse_clause_json(filepath: str) -> list[Clause]:
    """
    解析 clause_list JSON 文件，返回清洗后的条款列表。

    输入 JSON 结构示例::

        {
          "response": {
            "clause_list": [
              {"number": "1", "content": "...", "level": 1, "path": "", "full_name": "..."},
              ...
            ]
          }
        }

    处理逻辑：
    - 从 response.clause_list 提取条款数组
    - 将 content 字段中的 HTML（Quill 编辑器产出的表格等）转为 Markdown
    - content 中的编号前缀（如 "1.2.1.  "）已在原始数据中存在，保持原样
    - path 字段自动补全：原始 path 只含直接父级编号，自动构建完整祖先路径
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)

    clause_list: list[dict] = data.get('response', {}).get('clause_list', [])
    if not clause_list:
        raise ValueError(f"JSON 文件中未找到 response.clause_list: {filepath}")

    number_set: dict[str, dict] = {raw.get('number', ''): raw for raw in clause_list}

    result: list[Clause] = []
    for raw in clause_list:
        number = raw.get('number', '')
        if number == '前言':
            number = '0'
        raw_path = raw.get('path', '')

        if raw_path:
            path = _build_full_path(number, number_set)
        else:
            path = ''

        clause: Clause = {
            'number': number,
            'content': _convert_html_to_markdown(raw.get('content', '')),
            'level': raw.get('level', 0),
            'path': path,
            'full_name': raw.get('full_name', ''),
        }
        result.append(clause)

    logger.info(f"从 JSON 解析到 {len(result)} 个条款（已完成 HTML→Markdown 转换与路径补全）")
    return result


def parse_document(filepath: str) -> list[ParsedLine]:
    """统一文档解析入口，根据扩展名分派到对应解析器"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        logger.info(f"解析 docx 文件: {filepath}")
        return parse_docx(filepath)
    elif ext == ".txt":
        logger.info(f"解析 txt 文件: {filepath}")
        return parse_txt(filepath)
    else:
        raise ValueError(f"不支持的文件类型: {ext}，仅支持 .docx 和 .txt")


DEFAULT_API_URL = "http://10.199.0.40:8080/kg-platform/api/clauses/list"
DEFAULT_CLAUSE_API_URL = "http://10.199.0.40:8080/kg-platform/api/clauses/get"


def _fetch_single_clause(policy_id: str, clause_id: str, api_url: str) -> dict | None:
    """按 (policyId, clauseId) 调用细粒度接口拉取单条 clause raw dict。

    返回结构与 fetch_api_clauses 中 raw 同形（包含 clauseId/clauseNumber/clauseName/clauseContent 等）。
    失败、网络异常、查无均返回 None（调用方据此标记 missing/error）。
    """
    import requests
    try:
        response = requests.post(api_url, json={"policyId": policy_id, "clauseId": clause_id})
        response.raise_for_status()
        result = response.json()
    except Exception as e:
        logger.warning(f"_fetch_single_clause 网络/解析异常 policyId={policy_id} clauseId={clause_id}: {e}")
        return None

    if not result.get('success'):
        logger.warning(f"_fetch_single_clause 接口返回失败 policyId={policy_id} clauseId={clause_id}: {result.get('message')}")
        return None

    data = result.get('data')
    if data is None:
        return None
    if isinstance(data, list):
        if not data:
            return None
        for item in data:
            if isinstance(item, dict) and item.get('clauseId') == clause_id:
                return item
        first = data[0]
        return first if isinstance(first, dict) else None
    if isinstance(data, dict):
        if 'clause' in data and isinstance(data['clause'], dict):
            return data['clause']
        if 'clauses' in data and isinstance(data['clauses'], list) and data['clauses']:
            for item in data['clauses']:
                if isinstance(item, dict) and item.get('clauseId') == clause_id:
                    return item
            return data['clauses'][0] if isinstance(data['clauses'][0], dict) else None
        return data
    return None


def _fetch_policy_clauses_raw(policy_id: str, api_url: str) -> list[dict]:
    """按 policyId 调用全文接口拉取整篇 policy 的 raw clause 列表。

    与 fetch_api_clauses 内部使用的接口同源，但仅返回原始 clause 列表，不做 number / path 加工。
    异常或为空时返回 []，调用方据此降级处理。
    """
    import requests
    try:
        response = requests.post(api_url, json={"policyId": policy_id})
        response.raise_for_status()
        result = response.json()
    except Exception as e:
        logger.warning(f"_fetch_policy_clauses_raw 网络/解析异常 policyId={policy_id}: {e}")
        return []

    if not result.get('success'):
        logger.warning(f"_fetch_policy_clauses_raw 接口返回失败 policyId={policy_id}: {result.get('message')}")
        return []

    return result.get('data', {}).get('clauses', []) or []


def resolve_references(
    refs: list[dict],
    clause_cache: dict[tuple[str, str], dict | None],
    policy_cache: dict[str, list[str]],
    visited: set[tuple[str, str]],
    single_url: str,
    bulk_url: str,
) -> None:
    """就地填充 refs[i]['resolvedClauses']，递归展开嵌套引用，遇循环 / 查无 / 异常停止该分支。

    节点统一字段: policyId / clauseId / highlightedContent / resolvedClauses / cycle
    可选标记: missing(bool) / error(str)

    顶层 ref 保留原始 clauseId 串（含逗号或空串）与 highlightedContent；
    其 resolvedClauses 内为"扁平"的 clause 节点（每项 clauseId 都是单 id，highlightedContent 恒空），
    深层节点的 resolvedClauses 同样是扁平 clause 节点（不再嵌套 reference 对象）。

    见 plan 文档说明：
      - clauseId 空串 → 走 bulk 接口拉整篇，展开为 N 个单 id 子节点
      - 含逗号 → 拆分逐个走单条
      - 单 id → 走单条
      - 循环（在 visited 中）→ 节点 cycle:true，停止下钻
      - 缓存命中（不在 visited）→ 复用 raw 数据但仍递归
      - 接口失败/查无 → 节点 missing/error 并停止
    """
    for ref in refs:
        policy_id = ref.get('policyId', '')
        raw_clause_id = ref.get('clauseId', '')
        ref['resolvedClauses'] = _expand_to_clause_nodes(
            policy_id, raw_clause_id,
            clause_cache, policy_cache, visited, single_url, bulk_url,
        )


def _expand_to_clause_nodes(
    policy_id: str,
    raw_clause_id: str,
    clause_cache: dict[tuple[str, str], dict | None],
    policy_cache: dict[str, list[str]],
    visited: set[tuple[str, str]],
    single_url: str,
    bulk_url: str,
) -> list[dict]:
    """把 (policyId, raw_clauseId) 展开成扁平的 clause 节点列表，每个节点本身可再含 resolvedClauses。

    raw_clauseId:
      - "" 整篇 → 该 policy 全部 clauseId
      - "id1,id2" → 拆分
      - 单 id → 单元素列表
    """
    if not policy_id:
        return []

    if raw_clause_id == '':
        child_ids = _expand_whole_policy(policy_id, clause_cache, policy_cache, bulk_url)
    elif ',' in raw_clause_id:
        child_ids = [cid.strip() for cid in raw_clause_id.split(',') if cid.strip()]
    else:
        child_ids = [raw_clause_id]

    return [
        _resolve_single_clause(
            policy_id, cid, clause_cache, policy_cache, visited, single_url, bulk_url,
        )
        for cid in child_ids
    ]


def _expand_whole_policy(
    policy_id: str,
    clause_cache: dict[tuple[str, str], dict | None],
    policy_cache: dict[str, list[str]],
    bulk_url: str,
) -> list[str]:
    """整篇引用：返回该 policy 的全部 clauseId 列表（按 API 原顺序），首次访问时拉 bulk 并回填两个 cache。"""
    if policy_id in policy_cache:
        return policy_cache[policy_id]

    raw_list = _fetch_policy_clauses_raw(policy_id, bulk_url)
    ids: list[str] = []
    for raw in raw_list:
        cid = raw.get('clauseId', '')
        if not cid:
            continue
        ids.append(cid)
        clause_cache.setdefault((policy_id, cid), raw)
    policy_cache[policy_id] = ids
    if not ids:
        logger.warning(f"_expand_whole_policy 未能拉到 policyId={policy_id} 的任何 clause")
    return ids


def _resolve_single_clause(
    policy_id: str,
    clause_id: str,
    clause_cache: dict[tuple[str, str], dict | None],
    policy_cache: dict[str, list[str]],
    visited: set[tuple[str, str]],
    single_url: str,
    bulk_url: str,
) -> dict:
    """对单个 (policyId, clauseId) 解析为一个轻量 clause 节点。

    决策顺序: visited(循环检测) → clause_cache(请求级 memo) → 网络。
    返回节点的 resolvedClauses 是"扁平 clause 节点"列表（不嵌套 reference 对象）。
    """
    node: dict = {
        'policyId': policy_id,
        'clauseId': clause_id,
        'highlightedContent': '',
        'resolvedClauses': [],
        'cycle': False,
    }
    key = (policy_id, clause_id)

    if key in visited:
        node['cycle'] = True
        return node

    if key in clause_cache:
        raw = clause_cache[key]
    else:
        raw = _fetch_single_clause(policy_id, clause_id, single_url)
        clause_cache[key] = raw

    if raw is None:
        node['missing'] = True
        return node

    inner_html = raw.get('clauseContent', '') or ''
    inner_refs = extract_clause_references(inner_html)
    if not inner_refs:
        return node

    visited.add(key)
    try:
        flat_children: list[dict] = []
        for ref in inner_refs:
            children = _expand_to_clause_nodes(
                ref.get('policyId', ''), ref.get('clauseId', ''),
                clause_cache, policy_cache, visited, single_url, bulk_url,
            )
            flat_children.extend(children)
        node['resolvedClauses'] = flat_children
    finally:
        visited.discard(key)
    return node


def _extract_version_from_clause_name(clause_name: str) -> str:
    """从 clauseName 中提取括号内的版本号，如 '（20260413110049）' → '20260413110049'"""
    m = re.search(r'[（(](\d+)[）)]', clause_name)
    return m.group(1) if m else ''


def _extract_policy_name(clause_name: str) -> str:
    """从 clauseName 中提取策略名称（去掉 'clauseNumber_' 前缀和版本号后缀）"""
    name = re.sub(r'^[^_]+_', '', clause_name, count=1)
    name = re.sub(r'[（(]\d+[）)]', '', name).strip()
    return name


def fetch_api_clauses(
    policy_id: str,
    api_url: str = DEFAULT_API_URL,
    single_clause_api_url: str = DEFAULT_CLAUSE_API_URL,
) -> tuple[list[Clause], str, str, dict[str, dict]]:
    """
    从 API 接口获取条款数据并转换为 Clause 列表。

    参数:
        policy_id: 策略 ID
        api_url: 全文（按 policyId 拉整篇）接口地址
        single_clause_api_url: 细粒度（按 policyId+clauseId 拉单条）接口地址，用于 references 递归展开

    返回:
        (clauses, policy_name, version, raw_clause_map)
        - clauses: 转换后的 Clause 列表（与 parse_clause_json 格式一致）
        - policy_name: 从 clauseName 提取的策略名称
        - version: 从 clauseName 提取的版本号（用作目录时间戳后缀）
        - raw_clause_map: number → 原始 API 条款数据的映射（用于保存 clause.json，
          每个 raw 已被注入 references 字段：clauseContent 中 ql-reference 解析 + 递归展开结果）
    """
    import requests

    logger.info(f"从 API 获取条款数据: policyId={policy_id}, url={api_url}")
    response = requests.post(api_url, json={"policyId": policy_id}, timeout=60)
    response.raise_for_status()

    result = response.json()
    data_obj = result.get('data') or {}
    data_keys = list(data_obj.keys()) if isinstance(data_obj, dict) else type(data_obj).__name__
    api_clauses: list[dict] = data_obj.get('clauses', []) if isinstance(data_obj, dict) else []
    logger.info(
        f"上游响应概要: policyId={policy_id}, http={response.status_code}, "
        f"success={result.get('success')}, code={result.get('code')}, "
        f"message={result.get('message')!r}, data.keys={data_keys}, "
        f"clauses_count={len(api_clauses)}"
    )

    if not result.get('success'):
        raise ValueError(f"API 返回失败: {result.get('message', '未知错误')}")

    if not api_clauses:
        preview = json.dumps(result, ensure_ascii=False)[:500]
        raise ValueError(
            f"API 未返回有效条款数据，policyId: {policy_id}（上游响应预览: {preview}）"
        )

    policy_name = ''
    version = ''
    for c in api_clauses:
        cn = c.get('clauseName', '')
        v = _extract_version_from_clause_name(cn)
        if v:
            version = v
            if not policy_name:
                policy_name = _extract_policy_name(cn)
            break

    if not policy_name:
        policy_name = policy_id

    number_set: dict[str, dict] = {}
    raw_clause_map: dict[str, dict] = {}
    clauses: list[Clause] = []

    for raw in api_clauses:
        number = raw.get('clauseNumber', '')
        if number == '前言':
            number = '0'

        number_set[number] = raw
        raw_clause_map[number] = raw

        raw['references'] = extract_clause_references(raw.get('clauseContent', ''))

        search_labels = raw.get('searchLabels', [])
        full_name = search_labels[0] if search_labels else ''

        content = _convert_html_to_markdown(raw.get('clauseContent', ''))
        level = int(raw.get('level', 0))

        clause: Clause = {
            'number': number,
            'content': content,
            'level': level,
            'path': '',
            'full_name': full_name,
        }
        clauses.append(clause)

    for clause in clauses:
        number = clause['number']
        parent_num = _derive_parent_number(number)
        if parent_num and parent_num in number_set:
            clause['path'] = _build_full_path(number, number_set)

    # 递归展开 references 关系图（轻量节点，仅含 policyId/clauseId/highlightedContent/resolvedClauses/cycle）。
    # cache 仅在本次抽取生命周期内有效，不持久化、不跨任务复用。
    clause_cache: dict[tuple[str, str], dict | None] = {}
    policy_cache: dict[str, list[str]] = {}
    self_ids: list[str] = []
    for raw in api_clauses:
        cid = raw.get('clauseId', '')
        if cid:
            clause_cache[(policy_id, cid)] = raw
            self_ids.append(cid)
    policy_cache[policy_id] = self_ids

    for raw in api_clauses:
        cid = raw.get('clauseId', '')
        visited: set[tuple[str, str]] = {(policy_id, cid)} if cid else {(policy_id, '')}
        resolve_references(
            raw['references'], clause_cache, policy_cache, visited,
            single_clause_api_url, api_url,
        )

    logger.info(f"从 API 获取到 {len(clauses)} 个条款（policyId: {policy_id}，版本: {version}）")
    return clauses, policy_name, version, raw_clause_map
